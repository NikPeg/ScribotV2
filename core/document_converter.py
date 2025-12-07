"""
Модуль для конвертации документов между различными форматами.
"""

import asyncio
import contextlib
import logging
import os
import re

import qrcode
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import parse_xml
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

# Константы
MIN_PDF_SIZE_BYTES = 1000  # Минимальный размер PDF файла (1KB)
MAX_TOC_TITLE_LENGTH = 30  # Максимальная длина заголовка TOC для поиска
MIN_TEXT_LENGTH_FOR_SECTION = 10  # Минимальная длина текста для определения секции
MIN_TEXT_LENGTH_FOR_CHAPTER = 20  # Минимальная длина текста для определения главы
MIN_TEXT_LENGTH_FOR_PARAGRAPH = 5  # Минимальная длина текста для определения параграфа
MAX_TOC_LINES = 5  # Максимальное количество строк в TOC
MAX_HEADING_LENGTH = 100  # Максимальная длина заголовка
MIN_CONTENT_LENGTH = 50  # Минимальная длина контента после заголовка
MAX_SEARCH_RANGE = 30  # Максимальный диапазон поиска после элемента

# Логгер для модуля
logger = logging.getLogger(__name__)


async def compile_latex_to_pdf(tex_content: str, output_dir: str, filename: str) -> tuple[bool, str]:
    """
    Асинхронно компилирует LaTeX в PDF.
    Запускает pdflatex дважды для корректной генерации содержания, ссылок и библиографии.
    
    Args:
        tex_content: Содержимое LaTeX файла
        output_dir: Директория для выходных файлов
        filename: Имя файла без расширения
    
    Returns:
        Tuple[bool, str]: (успех, путь_к_файлу_или_ошибка)
    """
    tex_file = os.path.join(output_dir, f"{filename}.tex")
    pdf_file = os.path.join(output_dir, f"{filename}.pdf")
    
    try:
        # Записываем tex файл
        with open(tex_file, 'w', encoding='utf-8') as f:
            f.write(tex_content)
        
        # Первый проход pdflatex (генерирует .aux файлы)
        process1 = await asyncio.create_subprocess_exec(
            'pdflatex',
            '-interaction=nonstopmode',
            '-output-directory', output_dir,
            tex_file,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
            cwd=output_dir
        )
        
        stdout1, stderr1 = await process1.communicate()
        
        # Второй проход pdflatex (использует .aux для содержания и ссылок)
        process2 = await asyncio.create_subprocess_exec(
            'pdflatex',
            '-interaction=nonstopmode',
            '-output-directory', output_dir,
            tex_file,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
            cwd=output_dir
        )
        
        stdout2, stderr2 = await process2.communicate()
        
        # Проверяем результат: главное - наличие PDF файла
        # pdflatex может возвращать ненулевой код даже при успешной компиляции (warnings)
        if os.path.exists(pdf_file):
            # Проверяем размер файла - если он слишком маленький, возможно компиляция не удалась
            file_size = os.path.getsize(pdf_file)
            if file_size > MIN_PDF_SIZE_BYTES:
                return True, pdf_file
        
        # Если PDF не создан или слишком маленький - это реальная ошибка
        # Собираем полный текст ошибки без обрезки
        stdout1_text = stdout1.decode('utf-8', errors='ignore')
        stdout2_text = stdout2.decode('utf-8', errors='ignore')
        stderr1_text = stderr1.decode('utf-8', errors='ignore')
        stderr2_text = stderr2.decode('utf-8', errors='ignore')
        
        error_msg = f"LaTeX compilation failed. Return code: {process2.returncode}\n"
        if not os.path.exists(pdf_file):
            error_msg += "PDF file was not created.\n"
        else:
            error_msg += f"PDF file exists but is too small ({os.path.getsize(pdf_file)} bytes).\n"
        error_msg += f"\n=== First pass stdout ===\n{stdout1_text}\n\n"
        error_msg += f"=== First pass stderr ===\n{stderr1_text}\n\n"
        error_msg += f"=== Second pass stdout ===\n{stdout2_text}\n\n"
        error_msg += f"=== Second pass stderr ===\n{stderr2_text}"
        return False, error_msg
            
    except Exception as e:
        return False, f"Exception during LaTeX compilation: {e!s}"


async def convert_pdf_to_docx(pdf_path: str, output_dir: str, filename: str) -> tuple[bool, str]:  # noqa: PLR0912, PLR0915
    """
    Конвертирует PDF в DOCX используя LibreOffice через промежуточный формат ODT.
    LibreOffice не может напрямую конвертировать PDF в DOCX, поэтому используем ODT как промежуточный формат.
    
    Args:
        pdf_path: Путь к PDF файлу
        output_dir: Директория для выходных файлов
        filename: Имя файла без расширения
    
    Returns:
        Tuple[bool, str]: (успех, путь_к_файлу_или_ошибка)
    """
    docx_file = os.path.join(output_dir, f"{filename}.docx")
    
    # Проверяем существование PDF файла
    if not os.path.exists(pdf_path):
        error_msg = f"PDF файл не найден: {pdf_path}"
        logger.error(error_msg)
        return False, error_msg
    
    pdf_size = os.path.getsize(pdf_path)
    logger.info(f"Начинаю конвертацию PDF в DOCX через ODT: {pdf_path} (размер: {pdf_size} байт)")
    
    libreoffice_commands = [
        'libreoffice',  # Linux/Windows в PATH
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS стандартная установка
        '/usr/bin/libreoffice',  # Linux системная установка
        'soffice'  # Альтернативное имя
    ]
    
    last_error = None
    
    for cmd in libreoffice_commands:
        try:
            logger.debug(f"Проверяю доступность команды: {cmd}")
            # Проверяем доступность команды
            check_process = await asyncio.create_subprocess_exec(
                cmd, '--version',
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            _check_stdout, check_stderr = await check_process.communicate()
            
            if check_process.returncode == 0:
                logger.info(f"LibreOffice найден: {cmd}")
                
                # Шаг 1: Конвертируем PDF в ODT (LibreOffice может это делать)
                pdf_basename = os.path.basename(pdf_path)
                pdf_name_without_ext = os.path.splitext(pdf_basename)[0]
                odt_file = os.path.join(output_dir, f"{pdf_name_without_ext}.odt")
                
                logger.debug(f"Шаг 1: Конвертация PDF в ODT: {cmd} --headless --convert-to odt --outdir {output_dir} {pdf_path}")
                process_odt = await asyncio.create_subprocess_exec(
                    cmd,
                    '--headless',
                    '--convert-to', 'odt',
                    '--outdir', output_dir,
                    pdf_path,
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE
                )
                
                stdout_odt, stderr_odt = await process_odt.communicate()
                _stdout_odt_text = stdout_odt.decode('utf-8', errors='ignore') if stdout_odt else ""
                stderr_odt_text = stderr_odt.decode('utf-8', errors='ignore') if stderr_odt else ""
                
                logger.debug(f"PDF->ODT завершился с кодом: {process_odt.returncode}")
                if stderr_odt_text:
                    logger.debug(f"PDF->ODT stderr: {stderr_odt_text[:500]}")
                
                if process_odt.returncode != 0 or not os.path.exists(odt_file):
                    error_msg = (
                        f"Не удалось конвертировать PDF в ODT. "
                        f"Код возврата: {process_odt.returncode}, "
                        f"Файл существует: {os.path.exists(odt_file)}, "
                        f"stderr: {stderr_odt_text[:500]}"
                    )
                    logger.warning(error_msg)
                    last_error = error_msg
                    continue
                
                logger.info(f"ODT файл создан: {odt_file}")
                
                # Шаг 2: Конвертируем ODT в DOCX
                logger.debug(f"Шаг 2: Конвертация ODT в DOCX: {cmd} --headless --convert-to docx --outdir {output_dir} {odt_file}")
                process_docx = await asyncio.create_subprocess_exec(
                    cmd,
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', output_dir,
                    odt_file,
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE
                )
                
                stdout_docx, stderr_docx = await process_docx.communicate()
                stdout_docx_text = stdout_docx.decode('utf-8', errors='ignore') if stdout_docx else ""
                stderr_docx_text = stderr_docx.decode('utf-8', errors='ignore') if stderr_docx else ""
                
                logger.debug(f"ODT->DOCX завершился с кодом: {process_docx.returncode}")
                if stderr_docx_text:
                    logger.debug(f"ODT->DOCX stderr: {stderr_docx_text[:500]}")
                
                # LibreOffice создает файл с именем исходного ODT, но с расширением .docx
                generated_docx = os.path.join(output_dir, f"{pdf_name_without_ext}.docx")
                
                logger.debug(f"Ожидаемый файл: {generated_docx}, существует: {os.path.exists(generated_docx)}")
                logger.debug(f"Целевой файл: {docx_file}, существует: {os.path.exists(docx_file)}")
                
                # Удаляем промежуточный ODT файл
                with contextlib.suppress(OSError):
                    os.remove(odt_file)
                    logger.debug(f"Промежуточный ODT файл удален: {odt_file}")
                
                if process_docx.returncode == 0 and os.path.exists(generated_docx):
                    # Переименовываем в нужное имя
                    if generated_docx != docx_file:
                        try:
                            os.rename(generated_docx, docx_file)
                            logger.info(f"Файл переименован: {generated_docx} -> {docx_file}")
                        except OSError as e:
                            logger.warning(f"Не удалось переименовать файл: {e}")
                            # Пробуем использовать существующий файл
                            docx_file = generated_docx
                    
                    if os.path.exists(docx_file):
                        file_size = os.path.getsize(docx_file)
                        logger.info(f"DOCX файл успешно создан: {docx_file} (размер: {file_size} байт)")
                        return True, docx_file
                    error_msg = f"Файл {docx_file} не существует после переименования"
                    logger.error(error_msg)
                    last_error = error_msg
                else:
                    error_msg = (
                        f"LibreOffice конвертация ODT->DOCX не удалась. "
                        f"Код возврата: {process_docx.returncode}, "
                        f"Файл существует: {os.path.exists(generated_docx)}, "
                        f"stdout: {stdout_docx_text[:200]}, "
                        f"stderr: {stderr_docx_text[:200]}"
                    )
                    logger.error(error_msg)
                    last_error = error_msg
            else:
                logger.debug(f"Команда {cmd} недоступна (код возврата: {check_process.returncode})")
                check_stderr_text = check_stderr.decode('utf-8', errors='ignore') if check_stderr else ""
                last_error = f"Команда {cmd} недоступна: {check_stderr_text[:200]}"
                    
        except FileNotFoundError:
            logger.debug(f"Команда {cmd} не найдена")
            last_error = f"Команда {cmd} не найдена в PATH"
            continue
        except Exception as e:
            logger.error(f"Ошибка при попытке использовать {cmd}: {e}", exc_info=True)
            last_error = f"Ошибка при использовании {cmd}: {e!s}"
            continue
    
    error_msg = f"LibreOffice не найден или не может конвертировать PDF в DOCX. Последняя ошибка: {last_error}"
    logger.error(error_msg)
    return False, error_msg


async def convert_tex_to_docx(tex_content: str, output_dir: str, filename: str) -> tuple[bool, str]:
    """
    Конвертирует TEX в DOCX.
    Сначала пробует pandoc для прямой конвертации (наиболее надежный способ).
    Если pandoc не доступен, пробует LibreOffice через промежуточный PDF.
    
    Args:
        tex_content: Содержимое LaTeX файла
        output_dir: Директория для выходных файлов
        filename: Имя файла без расширения
    
    Returns:
        Tuple[bool, str]: (успех, путь_к_файлу_или_ошибка)
    """
    logger.info(f"Начинаю конвертацию TEX в DOCX для файла: {filename}")
    
    # Сначала пробуем pandoc для прямой конвертации TEX -> DOCX (наиболее надежный способ)
    logger.debug("Шаг 1: Пробую прямую конвертацию через pandoc")
    success, result = await _convert_tex_to_docx_direct(tex_content, output_dir, filename)
    if success:
        logger.info(f"DOCX успешно создан через pandoc: {result}")
        return True, result
    
    logger.warning("Pandoc не смог конвертировать, пробую альтернативные методы")
    
    # Если pandoc не сработал, пробуем через LibreOffice напрямую из TEX
    logger.debug("Шаг 2: Пробую LibreOffice напрямую из TEX")
    success, result = await _convert_via_libreoffice(tex_content, output_dir, filename)
    if success:
        logger.info(f"DOCX успешно создан через LibreOffice: {result}")
        return True, result
    
    # В крайнем случае пробуем через PDF (но это может не работать, так как LibreOffice не конвертирует PDF в ODT)
    logger.debug("Шаг 3: Пробую через промежуточный PDF")
    success, pdf_path = await compile_latex_to_pdf(tex_content, output_dir, filename)
    if success:
        logger.info(f"PDF успешно скомпилирован: {pdf_path}")
        # Пробуем конвертировать PDF в DOCX (может не работать)
        return await convert_pdf_to_docx(pdf_path, output_dir, filename)
    
    # Если ничего не сработало, возвращаем ошибку
    error_msg = "Не удалось конвертировать TEX в DOCX ни одним из доступных методов"
    logger.error(error_msg)
    return False, error_msg


async def _convert_tex_to_docx_direct(tex_content: str, output_dir: str, filename: str) -> tuple[bool, str]:
    """
    Прямая конвертация TEX в DOCX через pandoc.
    Это основной и наиболее надежный метод конвертации.
    
    Args:
        tex_content: Содержимое LaTeX файла
        output_dir: Директория для выходных файлов
        filename: Имя файла без расширения
    
    Returns:
        Tuple[bool, str]: (успех, путь_к_файлу_или_ошибка)
    """
    logger.info("Пробую прямую конвертацию TEX в DOCX через pandoc")
    docx_file = os.path.join(output_dir, f"{filename}.docx")
    
    # Сначала пробуем pandoc с улучшенными параметрами
    try:
        # Используем оригинальный LaTeX без модификаций, чтобы pandoc мог создать TOC
        # Pandoc с --toc создаст TOC как SDT элемент в начале документа
        # Затем мы программно переместим его после титульной страницы
        # Обрабатываем только \newpage, чтобы убрать "ewpage" из результата
        modified_tex = re.sub(r'\\newpage\s*', '\n\n', tex_content)
        modified_tex = re.sub(r'\n\s*\n\s*\n+', '\n\n', modified_tex)
        
        # Создаем временный tex файл
        tex_file = os.path.join(output_dir, f"{filename}_temp.tex")
        logger.debug(f"Создаю временный TEX файл: {tex_file}")
        with open(tex_file, 'w', encoding='utf-8') as f:
            f.write(modified_tex)
        
        # Используем --toc для генерации оглавления
        # Pandoc разместит TOC в начале, но мы модифицировали LaTeX так,
        # чтобы титульная страница была отделена, и TOC будет после нее
        logger.debug(f"Запускаю pandoc: pandoc {tex_file} -o {docx_file}")
        pandoc_process = await asyncio.create_subprocess_exec(
            'pandoc',
            tex_file,
            '-o', docx_file,
            '--from=latex',
            '--to=docx',
            '--toc',  # Генерировать оглавление
            '--toc-depth=3',  # Глубина оглавления
            '--wrap=none',  # Не переносить строки
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE
        )
        
        stdout, stderr = await pandoc_process.communicate()
        stdout_text = stdout.decode('utf-8', errors='ignore') if stdout else ""
        stderr_text = stderr.decode('utf-8', errors='ignore') if stderr else ""
        
        logger.debug(f"Pandoc завершился с кодом: {pandoc_process.returncode}")
        if stdout_text:
            logger.debug(f"Pandoc stdout: {stdout_text[:500]}")
        if stderr_text:
            logger.debug(f"Pandoc stderr: {stderr_text[:500]}")
        
        if pandoc_process.returncode == 0 and os.path.exists(docx_file):
            # Перемещаем TOC после титульной страницы
            try:
                _move_toc_after_title_page(docx_file)
                logger.info("TOC успешно перемещен после титульной страницы")
            except Exception as e:
                logger.warning(f"Не удалось переместить TOC: {e}. Оставляем TOC в начале документа.")
            
            # Добавляем разрывы страниц в нужных местах
            try:
                _add_page_breaks_to_docx(docx_file)
                logger.info("Разрывы страниц успешно добавлены")
            except Exception as e:
                logger.warning(f"Не удалось добавить разрывы страниц: {e}")
            
            file_size = os.path.getsize(docx_file)
            logger.info(f"DOCX успешно создан через pandoc: {docx_file} (размер: {file_size} байт)")
            # Удаляем временный файл
            with contextlib.suppress(OSError):
                os.remove(tex_file)
            return True, docx_file
        error_msg = (
            f"Pandoc конвертация не удалась. "
            f"Код возврата: {pandoc_process.returncode}, "
            f"Файл существует: {os.path.exists(docx_file)}, "
            f"stderr: {stderr_text[:500]}"
        )
        logger.warning(error_msg)
        # Удаляем временный файл
        with contextlib.suppress(OSError):
            os.remove(tex_file)
        return False, error_msg
            
    except FileNotFoundError:
        logger.warning("Pandoc не найден в PATH")
        return False, "Pandoc не найден в PATH"
    except Exception as e:
        logger.error(f"Ошибка при использовании pandoc: {e}", exc_info=True)
        # Удаляем временный файл
        with contextlib.suppress(OSError):
            if 'tex_file' in locals():
                os.remove(tex_file)
        return False, f"Ошибка при использовании pandoc: {e!s}"


async def _convert_via_libreoffice(tex_content: str, output_dir: str, filename: str) -> tuple[bool, str]:
    """
    Конвертирует через LibreOffice как резервный метод.
    
    Args:
        tex_content: Содержимое LaTeX файла
        output_dir: Директория для выходных файлов
        filename: Имя файла без расширения
    
    Returns:
        Tuple[bool, str]: (успех, путь_к_файлу_или_ошибка)
    """
    docx_file = os.path.join(output_dir, f"{filename}.docx")
    
    libreoffice_commands = [
        'libreoffice',  # Linux/Windows в PATH
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',  # macOS стандартная установка
        '/usr/bin/libreoffice',  # Linux системная установка
        'soffice'  # Альтернативное имя
    ]
    
    for cmd in libreoffice_commands:
        try:
            # Проверяем доступность команды
            check_process = await asyncio.create_subprocess_exec(
                cmd, '--version',
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE
            )
            await check_process.communicate()
            
            if check_process.returncode == 0:
                # Создаем простой ODT файл из текста (без LaTeX команд)
                # Извлекаем только текстовое содержимое
                clean_text = _extract_text_from_latex(tex_content)
                
                # Создаем простой текстовый файл
                txt_file = os.path.join(output_dir, f"{filename}_temp.txt")
                with open(txt_file, 'w', encoding='utf-8') as f:
                    f.write(clean_text)
                
                # Конвертируем TXT в DOCX
                process = await asyncio.create_subprocess_exec(
                    cmd,
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', output_dir,
                    txt_file,
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE
                )
                
                _stdout, _stderr = await process.communicate()
                
                # Переименовываем результат
                txt_docx = os.path.join(output_dir, f"{filename}_temp.docx")
                if process.returncode == 0 and os.path.exists(txt_docx):
                    with contextlib.suppress(OSError):
                        os.rename(txt_docx, docx_file)
                        os.remove(txt_file)
                        return True, docx_file
                
                # Очищаем временные файлы
                with contextlib.suppress(OSError):
                    os.remove(txt_file)
                    if os.path.exists(txt_docx):
                        os.remove(txt_docx)
                    
        except Exception:
            continue
    
    return False, "Neither pandoc nor LibreOffice could convert to DOCX"


def _move_toc_after_title_page(docx_path: str) -> None:  # noqa: PLR0912, PLR0915
    """
    Перемещает оглавление (TOC) после титульной страницы в DOCX файле.
    Pandoc с --toc всегда размещает TOC в начале документа, поэтому
    мы программно перемещаем его после титульной страницы.
    
    Args:
        docx_path: Путь к DOCX файлу
    """
    try:
        doc = Document(docx_path)
        body = doc.element.body
        paragraphs = list(doc.paragraphs)
        
        logger.info(f"Всего параграфов в документе: {len(paragraphs)}")
        
        # Проверяем, есть ли SDT элемент (TOC) в начале документа
        toc_sdt = None
        if len(body) > 0:
            first_elem = body[0]
            # Проверяем, является ли первый элемент SDT (structured document tag - TOC)
            if 'sdt' in first_elem.tag.lower():
                toc_sdt = first_elem
                logger.info("Найден TOC как SDT элемент в начале документа")
        
        # Ищем начало TOC в параграфах - ищем текст "Table of Contents"
        toc_start_idx = None
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            text_lower = text.lower()
            # Ищем различные варианты заголовка TOC
            if ('table of contents' in text_lower or
                'содержание' in text_lower or
                'оглавление' in text_lower or
                (text and len(text) < MAX_TOC_TITLE_LENGTH and 'contents' in text_lower)):
                toc_start_idx = i
                logger.info(f"Найден TOC в параграфах на позиции {i}: '{text[:60]}'")
                break
        
        # Если TOC найден как SDT элемент, обрабатываем его отдельно
        if toc_sdt is not None:
            logger.info("Обрабатываю TOC как SDT элемент")
            # Находим титульную страницу в параграфах
            title_end_idx = None
            for i, para in enumerate(paragraphs):
                text = para.text.strip()
                if 'проверил:' in text.lower() or ('петров' in text.lower() and 'п.п' in text.lower()):
                    title_end_idx = i
                    for j in range(i + 1, min(i + 3, len(paragraphs))):
                        if paragraphs[j].text.strip():
                            title_end_idx = j
                        else:
                            break
                    break
            
            if title_end_idx is not None:
                # Находим элемент титульной страницы в body
                title_elem = None
                para_count = 0
                for elem in body:
                    if 'p' in elem.tag.lower():
                        if para_count == title_end_idx:
                            title_elem = elem
                            break
                        para_count += 1
                
                if title_elem is not None:
                    # Перемещаем SDT элемент после титульной страницы
                    parent = toc_sdt.getparent()
                    if parent is not None:
                        # Удаляем SDT из текущей позиции
                        parent.remove(toc_sdt)
                        # Вставляем после титульной страницы
                        parent.insert(parent.index(title_elem) + 1, toc_sdt)
                        doc.save(docx_path)
                        logger.info("SDT TOC успешно перемещен после титульной страницы")
                        return
            else:
                logger.warning("Не удалось найти титульную страницу для перемещения SDT TOC")
        
        if toc_start_idx is None:
            logger.warning("TOC не найден в документе - пропускаем перемещение")
            return
        
        # Находим конец TOC - ищем начало титульной страницы ("МИНИСТЕРСТВО")
        toc_end_idx = None
        title_start_idx = None
        for i in range(toc_start_idx + 1, len(paragraphs)):
            text = paragraphs[i].text.strip()
            text_lower = text.lower()
            # Ищем начало титульной страницы
            if 'министерство' in text_lower or 'российский государственный университет' in text_lower:
                toc_end_idx = i
                title_start_idx = i
                logger.info(f"Найден конец TOC и начало титульной страницы на позиции {i}")
                break
        
        if toc_end_idx is None:
            logger.warning("Не удалось найти конец TOC - пропускаем перемещение")
            return
        
        # Ищем конец титульной страницы (ищем "Проверил:" или "Петров П.П.")
        title_end_idx = None
        for i in range(title_start_idx, len(paragraphs)):
            text = paragraphs[i].text.strip()
            if 'проверил:' in text.lower() or ('петров' in text.lower() and 'п.п' in text.lower()):
                # Ищем последний параграф титульной страницы
                title_end_idx = i
                # Продолжаем искать еще 1-2 параграфа после "Проверил:"
                for j in range(i + 1, min(i + 3, len(paragraphs))):
                    next_text = paragraphs[j].text.strip()
                    if next_text:
                        title_end_idx = j
                    else:
                        break
                logger.info(f"Найден конец титульной страницы на позиции {title_end_idx}")
                break
        
        if title_end_idx is None:
            logger.warning("Не удалось найти конец титульной страницы - пропускаем перемещение")
            return
        
        logger.info(f"Перестраиваю документ: TOC ({toc_start_idx}-{toc_end_idx}), Титульная ({title_start_idx}-{title_end_idx})")
        
        # Создаем новый документ с правильным порядком: титульная страница → TOC → контент
        new_doc = Document()
        
        # Функция для копирования параграфа
        def copy_paragraph(src_para, dst_doc):
            new_para = dst_doc.add_paragraph()
            new_para.text = src_para.text
            if src_para.style:
                with contextlib.suppress(Exception):
                    new_para.style = src_para.style
            # Копируем форматирование runs
            for run in src_para.runs:
                new_run = new_para.add_run(run.text)
                with contextlib.suppress(Exception):
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.size:
                        new_run.font.size = run.font.size
        
        # 1. Копируем титульную страницу
        for i in range(title_start_idx, title_end_idx + 1):
            copy_paragraph(paragraphs[i], new_doc)
        
        # 2. Копируем TOC
        for i in range(toc_start_idx, toc_end_idx):
            copy_paragraph(paragraphs[i], new_doc)
        
        # 3. Копируем остальные параграфы (основной контент)
        # Пропускаем TOC и титульную страницу
        for i in range(len(paragraphs)):
            if toc_start_idx <= i < toc_end_idx:
                continue  # Пропускаем TOC (уже скопирован)
            if title_start_idx <= i <= title_end_idx:
                continue  # Пропускаем титульную страницу (уже скопирована)
            copy_paragraph(paragraphs[i], new_doc)
        
        # Сохраняем новый документ
        new_doc.save(docx_path)
        logger.info("TOC успешно перемещен после титульной страницы")
            
    except Exception as e:
        logger.error(f"Ошибка при перемещении TOC: {e}", exc_info=True)
        raise


def _prepare_tex_for_pandoc(tex_content: str) -> str:
    """
    Подготавливает LaTeX контент для конвертации через pandoc.
    Pandoc с --toc всегда размещает TOC в начале документа.
    После конвертации мы программно переместим TOC после титульной страницы.
    
    Args:
        tex_content: Исходный LaTeX контент
    
    Returns:
        Модифицированный LaTeX контент для pandoc
    """
    # Оставляем \tableofcontents - pandoc может его обработать, но создаст TOC в начале
    # Затем мы программно переместим его после титульной страницы
    # НЕ удаляем \tableofcontents, чтобы pandoc мог создать TOC
    
    # Обрабатываем \newpage - заменяем на двойной перенос строки
    result = re.sub(r'\\newpage\s*', '\n\n', tex_content)
    
    # Убираем лишние пустые строки
    return re.sub(r'\n\s*\n\s*\n+', '\n\n', result)


def _add_page_breaks_to_docx(docx_path: str) -> None:  # noqa: PLR0912, PLR0915
    """
    Добавляет разрывы страниц в DOCX файле в нужных местах:
    1. После титульной страницы (перед TOC)
    2. После TOC (перед первой главой)
    3. Перед каждой новой главой (section)
    
    Args:
        docx_path: Путь к DOCX файлу
    """
    try:
        doc = Document(docx_path)
        paragraphs = list(doc.paragraphs)
        
        logger.info(f"Добавляю разрывы страниц в документ с {len(paragraphs)} параграфами")
        
        # Выводим первые несколько параграфов для отладки
        for i, para in enumerate(paragraphs[:10]):
            logger.debug(f"Параграф {i}: '{para.text[:60]}...' (стиль: {para.style.name if para.style else 'None'})")
        
        # Находим позиции ключевых элементов
        title_end_idx = None
        toc_start_idx = None
        toc_end_idx = None
        
        # Находим конец титульной страницы
        # Титульная страница заканчивается на параграфе с "Проверил:" или "Петров П.П."
        for i, para in enumerate(paragraphs):
            text = para.text.strip()
            if 'проверил:' in text.lower() or ('петров' in text.lower() and 'п.п' in text.lower()):
                title_end_idx = i
                logger.info(f"Найден конец титульной страницы на позиции {i}: '{text[:60]}'")
                # Ищем последний параграф титульной страницы (может быть пустая строка после)
                for j in range(i + 1, min(i + 3, len(paragraphs))):
                    next_text = paragraphs[j].text.strip()
                    # Если следующий параграф не пустой и не является началом TOC или главы
                    if next_text and not any(marker in next_text.lower() for marker in ['table of contents', 'содержание', 'оглавление']):
                        # Проверяем, не является ли это заголовком главы
                        next_para = paragraphs[j]
                        if next_para.style and 'heading' in next_para.style.name.lower():
                            break
                        title_end_idx = j
                    else:
                        break
                break
        
        # Находим TOC (может быть в параграфах или как SDT элемент)
        body = doc.element.body
        
        # Сначала проверяем, есть ли SDT элемент (TOC)
        toc_sdt_idx = None
        for i, elem in enumerate(body):
            if 'sdt' in elem.tag.lower():
                toc_sdt_idx = i
                logger.info(f"Найден TOC как SDT элемент на позиции {i} в body")
                break
        
        # Ищем TOC в параграфах
        # TOC обычно содержит текст "Table of Contents" или "Содержание"
        for i, para in enumerate(paragraphs):
            text = para.text.strip().lower()
            # Проверяем, что это действительно TOC, а не заголовок главы
            # TOC обычно не является заголовком (Heading стиль)
            is_heading = para.style and 'heading' in para.style.name.lower()
            if (('table of contents' in text or 'содержание' in text or 'оглавление' in text) and not is_heading):
                toc_start_idx = i
                logger.info(f"Найден TOC в параграфах на позиции {i}")
                # Ищем конец TOC - следующую секцию или начало контента
                for j in range(i + 1, len(paragraphs)):
                    para_text = paragraphs[j].text.strip()
                    # Если нашли секцию (обычно начинается с номера или заголовка)
                    if (para_text and (
                        para_text.startswith('\\section') or
                        (len(para_text) > 0 and para_text[0].isdigit()) or
                        ('section' in para_text.lower() and len(para_text) < MAX_HEADING_LENGTH)
                    ) and j - i > MAX_TOC_LINES):  # TOC обычно занимает несколько строк
                        toc_end_idx = j
                        break
                # Если не нашли конец, используем следующую секцию после TOC
                if toc_end_idx is None:
                    for j in range(i + 1, min(i + MAX_SEARCH_RANGE, len(paragraphs))):
                        para_text = paragraphs[j].text.strip()
                        # Ищем начало первой главы (обычно это section или заголовок)
                        if (para_text and len(para_text) > MIN_TEXT_LENGTH_FOR_SECTION and
                            not any(marker in para_text.lower() for marker in ['table of contents', 'содержание', 'оглавление']) and
                            (len(para_text) > MIN_TEXT_LENGTH_FOR_CHAPTER or para_text[0].isdigit())):
                            toc_end_idx = j
                            break
                break
        
        # Если TOC найден как SDT элемент, определяем его позицию в параграфах
        if (toc_sdt_idx is not None and toc_start_idx is None and title_end_idx is not None):
            # Ищем позицию SDT элемента в параграфах
            # SDT элемент обычно находится после титульной страницы
            # TOC должен быть сразу после титульной страницы
            toc_start_idx = title_end_idx + 1
            # Ищем конец TOC - следующую секцию
            for j in range(toc_start_idx + 1, min(toc_start_idx + MAX_SEARCH_RANGE, len(paragraphs))):
                para_text = paragraphs[j].text.strip()
                para_j = paragraphs[j]
                if (para_text and len(para_text) > MIN_TEXT_LENGTH_FOR_SECTION and
                    not any(marker in para_text.lower() for marker in ['table of contents', 'содержание', 'оглавление']) and
                    (len(para_text) > MIN_TEXT_LENGTH_FOR_CHAPTER or para_text[0].isdigit() or
                     (para_j.style and 'heading' in para_j.style.name.lower()))):
                    toc_end_idx = j
                    break
        
        logger.info(f"Найдены позиции: title_end_idx={title_end_idx}, toc_start_idx={toc_start_idx}, toc_end_idx={toc_end_idx}, toc_sdt_idx={toc_sdt_idx}")
        
        # Добавляем разрыв страницы ПЕРЕД TOC (если TOC есть)
        if title_end_idx is not None and (toc_start_idx is not None or toc_sdt_idx is not None):
            # TOC найден, добавляем разрыв страницы ПЕРЕД TOC
            if toc_sdt_idx is not None:
                # TOC - SDT элемент, работаем с body напрямую
                # Находим SDT элемент в body
                sdt_elem = body[toc_sdt_idx]
                # Создаем параграф с разрывом страницы перед SDT элементом
                break_para_xml = '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:r><w:br w:type="page"/></w:r></w:p>'
                break_para_elem = parse_xml(break_para_xml)
                # Вставляем перед SDT элементом
                parent = sdt_elem.getparent()
                if parent is not None:
                    sdt_index = parent.index(sdt_elem)
                    parent.insert(sdt_index, break_para_elem)
                    logger.info(f"Добавлен разрыв страницы ПЕРЕД TOC (SDT элемент на позиции {toc_sdt_idx} в body)")
                    # Обновляем список параграфов
                    paragraphs = list(doc.paragraphs)
            elif toc_start_idx is not None:
                # TOC найден в параграфах
                target_para = paragraphs[toc_start_idx]
                # Проверяем, нет ли уже разрыва страницы перед TOC
                has_page_break = False
                if target_para.runs:
                    first_run = target_para.runs[0]
                    if hasattr(first_run, '_element'):
                        xml = first_run._element.xml
                        if 'w:br' in xml and 'w:type="page"' in xml:
                            has_page_break = True
                
                if not has_page_break:
                    # Добавляем разрыв страницы в начало TOC параграфа
                    if target_para.runs:
                        first_run = target_para.runs[0]
                        run_element = first_run._element
                        break_xml = '<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="page"/>'
                        break_element = parse_xml(break_xml)
                        run_element.insert(0, break_element)
                    else:
                        target_para.add_run().add_break(WD_BREAK.PAGE)
                    logger.info(f"Добавлен разрыв страницы ПЕРЕД TOC (позиция {toc_start_idx})")
                    paragraphs = list(doc.paragraphs)
        
        # Добавляем разрыв страницы после титульной страницы (если TOC нет, то перед первой главой)
        # Ищем первый непустой параграф после титульной страницы (TOC или заголовок главы)
        if title_end_idx is not None and toc_start_idx is None and toc_sdt_idx is None:
            # Ищем первый непустой параграф после титульной страницы (TOC или заголовок главы)
            target_idx = None
            for i in range(title_end_idx + 1, min(title_end_idx + 15, len(paragraphs))):
                para = paragraphs[i]
                text = para.text.strip()
                # Пропускаем пустые параграфы
                if not text or len(text) == 0:
                    continue
                
                text_lower = text.lower()
                is_toc = ('table of contents' in text_lower or 'содержание' in text_lower or 'оглавление' in text_lower)
                is_heading = para.style and 'heading' in para.style.name.lower()
                
                # Это либо TOC, либо заголовок главы - это наш целевой элемент
                if is_toc or is_heading:
                    target_idx = i
                    logger.info(f"Найден целевой элемент на позиции {i}: TOC={is_toc}, Heading={is_heading}, текст='{text[:50]}'")
                    break
            
            # Если не нашли явный TOC или заголовок, берем первый непустой параграф с достаточным количеством текста
            if target_idx is None:
                for i in range(title_end_idx + 1, min(title_end_idx + 15, len(paragraphs))):
                    para = paragraphs[i]
                    text = para.text.strip()
                    if text and len(text) > MIN_TEXT_LENGTH_FOR_PARAGRAPH:
                        target_idx = i
                        logger.info(f"Найден первый непустой параграф на позиции {i}: '{text[:50]}'")
                        break
            
            if target_idx is not None:
                target_para = paragraphs[target_idx]
                # Проверяем, нет ли уже разрыва страницы в начале целевого параграфа
                has_page_break = False
                if target_para.runs:
                    first_run = target_para.runs[0]
                    if hasattr(first_run, '_element'):
                        xml = first_run._element.xml
                        if 'w:br' in xml and 'w:type="page"' in xml:
                            has_page_break = True
                
                # Также проверяем предыдущий параграф
                if not has_page_break and target_idx > 0:
                    prev_para = paragraphs[target_idx - 1]
                    # Если предыдущий параграф пустой, проверяем его на разрыв страницы
                    if not prev_para.text.strip():
                        for run in prev_para.runs:
                            if hasattr(run, '_element'):
                                xml = run._element.xml
                                if 'w:br' in xml and 'w:type="page"' in xml:
                                    has_page_break = True
                                    break
                
                if not has_page_break:
                    # Добавляем разрыв страницы в начало целевого параграфа
                    # Используем XML напрямую для вставки разрыва страницы в начало первого run
                    if target_para.runs:
                        first_run = target_para.runs[0]
                        # Получаем XML первого run
                        run_element = first_run._element
                        # Создаем элемент разрыва страницы
                        break_xml = '<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="page"/>'
                        break_element = parse_xml(break_xml)
                        # Вставляем разрыв страницы в начало run (перед текстом)
                        run_element.insert(0, break_element)
                    else:
                        # Параграф без runs, добавляем run с разрывом
                        run = target_para.add_run()
                        run.add_break(WD_BREAK.PAGE)
                    target_text = target_para.text[:50].replace('\n', ' ')
                    logger.info(f"Добавлен разрыв страницы после титульной страницы (в начало параграфа {target_idx}: '{target_text}')")
                    # Обновляем список параграфов
                    paragraphs = list(doc.paragraphs)
        
        # Добавляем разрыв страницы после TOC (перед первой главой)
        # Если TOC найден, добавляем разрыв после него
        # Если TOC не найден, но есть титульная страница, добавляем разрыв перед первой главой
        if toc_end_idx is not None:
            # TOC найден, добавляем разрыв после него
            insert_idx = toc_end_idx
            # Обновляем список параграфов, так как могли добавить новые
            paragraphs = list(doc.paragraphs)
            if insert_idx < len(paragraphs):
                para = paragraphs[insert_idx]
                # Проверяем, нет ли уже разрыва страницы
                has_page_break = False
                if para.runs:
                    first_run = para.runs[0]
                    if hasattr(first_run, '_element'):
                        xml = first_run._element.xml
                        if 'w:br' in xml and 'w:type="page"' in xml:
                            has_page_break = True
                
                if not has_page_break:
                    # Создаем новый параграф с разрывом страницы перед первой главой
                    new_para = para.insert_paragraph_before()
                    new_para.add_run().add_break(WD_BREAK.PAGE)
                    logger.info(f"Добавлен разрыв страницы после TOC (перед позицией {insert_idx})")
        elif title_end_idx is not None and toc_start_idx is None and toc_sdt_idx is None:
            # TOC не найден, но есть титульная страница - ищем первую главу
            for j in range(title_end_idx + 1, min(title_end_idx + MAX_SEARCH_RANGE, len(paragraphs))):
                para_text = paragraphs[j].text.strip()
                para_j = paragraphs[j]
                if para_text and (
                    (para_j.style and 'heading' in para_j.style.name.lower()) or
                    para_text[0].isdigit() or
                    len(para_text) > MIN_TEXT_LENGTH_FOR_CHAPTER
                ):
                    # Нашли первую главу
                    # Обновляем список параграфов
                    paragraphs = list(doc.paragraphs)
                    if j < len(paragraphs):
                        para = paragraphs[j]
                        # Проверяем, нет ли уже разрыва страницы
                        has_page_break = False
                        if para.runs:
                            first_run = para.runs[0]
                            if hasattr(first_run, '_element'):
                                xml = first_run._element.xml
                                if 'w:br' in xml and 'w:type="page"' in xml:
                                    has_page_break = True
                        
                        if not has_page_break:
                            # Создаем новый параграф с разрывом страницы перед первой главой
                            new_para = para.insert_paragraph_before()
                            new_para.add_run().add_break(WD_BREAK.PAGE)
                            logger.info(f"Добавлен разрыв страницы перед первой главой (позиция {j})")
                    break
        
        # Находим все секции и добавляем разрыв страницы перед каждой новой секцией
        # Pandoc конвертирует \section в заголовки с определенными стилями
        # Ищем параграфы, которые являются заголовками (Heading 1, Heading 2, и т.д.)
        previous_section_idx = None
        
        for i, para in enumerate(paragraphs):
            # Пропускаем титульную страницу и TOC
            if title_end_idx is not None and i <= title_end_idx:
                continue
            if toc_start_idx is not None and toc_end_idx is not None and toc_start_idx <= i <= toc_end_idx:
                continue
            
            text = para.text.strip()
            if not text:
                continue
            
            # Проверяем, является ли параграф заголовком
            is_section = False
            
            # 1. Проверяем стиль параграфа (pandoc создает заголовки с определенными стилями)
            if para.style and para.style.name:
                style_name = para.style.name.lower()
                if 'heading' in style_name or 'заголовок' in style_name:
                    is_section = True
            
            # 2. Проверяем, начинается ли текст с номера секции (1., 2., и т.д.)
            if not is_section and text and (re.match(r'^\d+[.)]\s+[А-ЯЁ]', text) or re.match(r'^\d+[.)]\s+[A-Z]', text)):
                # Паттерн: число, точка/скобка, пробел, текст
                is_section = True
            
            # 3. Проверяем, является ли это коротким текстом, который может быть заголовком
            if (not is_section and text and len(text) < MAX_HEADING_LENGTH and text[0].isupper() and
                i + 1 < len(paragraphs)):
                next_text = paragraphs[i + 1].text.strip()
                # Если следующий параграф длинный (это контент главы), то текущий - заголовок
                if len(next_text) > MIN_CONTENT_LENGTH:
                    is_section = True
            
            if is_section:
                # Добавляем разрыв страницы перед секцией (кроме первой)
                if previous_section_idx is not None:
                    # Обновляем список параграфов, так как могли добавить новые
                    paragraphs = list(doc.paragraphs)
                    if i < len(paragraphs):
                        para = paragraphs[i]
                        # Проверяем, нет ли уже разрыва страницы
                        has_page_break = False
                        if para.runs:
                            first_run = para.runs[0]
                            if hasattr(first_run, '_element'):
                                xml = first_run._element.xml
                                if 'w:br' in xml and 'w:type="page"' in xml:
                                    has_page_break = True
                        
                        if not has_page_break:
                            # Добавляем разрыв страницы в начало параграфа секции (используем XML)
                            if para.runs:
                                first_run = para.runs[0]
                                # Получаем XML первого run
                                run_element = first_run._element
                                # Создаем элемент разрыва страницы
                                break_xml = '<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="page"/>'
                                break_element = parse_xml(break_xml)
                                # Вставляем разрыв страницы в начало run (перед текстом)
                                run_element.insert(0, break_element)
                            else:
                                # Параграф без runs, добавляем run с разрывом
                                para.add_run().add_break(WD_BREAK.PAGE)
                            logger.info(f"Добавлен разрыв страницы перед секцией на позиции {i}: '{text[:50]}'")
                previous_section_idx = i
        
        # Сохраняем документ
        doc.save(docx_path)
        logger.info("Разрывы страниц успешно добавлены в документ")
        
    except Exception as e:
        logger.error(f"Ошибка при добавлении разрывов страниц: {e}", exc_info=True)
        raise


def _extract_text_from_latex(tex_content: str) -> str:
    """
    Извлекает чистый текст из LaTeX, убирая команды форматирования.
    
    Args:
        tex_content: LaTeX содержимое
    
    Returns:
        Чистый текст
    """
    # Убираем LaTeX команды и оставляем только текст
    clean_text = re.sub(r'\\[a-zA-Z]+\{[^}]*\}', '', tex_content)
    clean_text = re.sub(r'\\[a-zA-Z]+', '', clean_text)
    clean_text = re.sub(r'\{[^}]*\}', '', clean_text)
    clean_text = re.sub(r'\\\\', '\n', clean_text)
    return re.sub(r'\n\s*\n', '\n\n', clean_text)


def _create_qr_code_image(payment_url: str, user_id: int, temp_dir: str) -> str:
    """
    Создает QR-код из ссылки на оплату и сохраняет его во временный файл.
    
    Args:
        payment_url: Ссылка на оплату
        user_id: ID пользователя (для уникального имени файла)
        temp_dir: Временная директория
    
    Returns:
        Путь к файлу с QR-кодом
    """
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(payment_url)
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="#220d8c", back_color="white")
    qr_path = os.path.join(temp_dir, f"qr_code_{user_id}.png")
    img.save(qr_path)
    
    return qr_path


def _create_qr_code_pdf_page(payment_url: str, user_id: int, temp_dir: str) -> str:
    """
    Создает PDF страницу с QR-кодом.
    
    Args:
        payment_url: Ссылка на оплату
        user_id: ID пользователя (для уникального имени файла)
        temp_dir: Временная директория
    
    Returns:
        Путь к PDF файлу с одной страницей
    """
    # Создаем QR-код
    qr_path = _create_qr_code_image(payment_url, user_id, temp_dir)
    
    # Создаем PDF страницу
    pdf_path = os.path.join(temp_dir, f"qr_page_{user_id}.pdf")
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    
    # Размер QR-кода - половина ширины страницы
    qr_size = width * 0.5
    
    # Позиция QR-кода по центру страницы
    qr_x = (width - qr_size) / 2
    qr_y = (height - qr_size) / 2
    
    # Вставляем QR-код
    c.drawImage(qr_path, qr_x, qr_y, width=qr_size, height=qr_size)
    
    c.save()
    
    return pdf_path


async def create_partial_pdf_with_qr(
    full_pdf_path: str,
    payment_url: str,
    user_id: int,
    temp_dir: str,
    output_filename: str
) -> tuple[bool, str]:
    """
    Создает частичный PDF: первая половина страниц из оригинала + страницы с QR-кодами.
    
    Args:
        full_pdf_path: Путь к полному PDF файлу
        payment_url: Ссылка на оплату
        user_id: ID пользователя
        temp_dir: Временная директория
        output_filename: Имя выходного файла без расширения
    
    Returns:
        Tuple[bool, str]: (успех, путь_к_файлу_или_ошибка)
    """
    try:
        # Читаем оригинальный PDF
        reader = PdfReader(full_pdf_path)
        total_pages = len(reader.pages)
        
        if total_pages == 0:
            return False, "PDF файл не содержит страниц"
        
        # Определяем количество страниц для первой половины
        half_pages = total_pages // 2
        qr_pages_count = total_pages - half_pages
        
        # Создаем новый PDF writer
        writer = PdfWriter()
        
        # Добавляем первую половину страниц из оригинала
        for i in range(half_pages):
            writer.add_page(reader.pages[i])
        
        # Создаем страницы с QR-кодами
        qr_page_path = _create_qr_code_pdf_page(payment_url, user_id, temp_dir)
        qr_reader = PdfReader(qr_page_path)
        qr_page = qr_reader.pages[0]
        
        # Добавляем страницы с QR-кодами
        for _ in range(qr_pages_count):
            writer.add_page(qr_page)
        
        # Сохраняем частичный PDF
        partial_pdf_path = os.path.join(temp_dir, f"{output_filename}_partial.pdf")
        with open(partial_pdf_path, 'wb') as output_file:
            writer.write(output_file)
        
        return True, partial_pdf_path
        
    except Exception as e:
        return False, f"Ошибка при создании частичного PDF: {e!s}"
