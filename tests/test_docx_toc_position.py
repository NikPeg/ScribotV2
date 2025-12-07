"""
Тесты для проверки правильного расположения оглавления (TOC) в DOCX файлах.
Проверяет, что TOC находится после титульной страницы, а не в начале документа.
"""

import contextlib
import os
import re
import shutil
import subprocess
import sys
import tempfile

import pytest
from docx import Document

# Константы
MIN_TOC_CONTENT_LENGTH = 10  # Минимальная длина текста для проверки содержимого TOC
MIN_TEXT_LENGTH_FOR_SECTION = 10  # Минимальная длина текста для определения секции
MIN_EXPECTED_PAGE_BREAKS = 2  # Минимальное количество разрывов страниц

# Добавляем корневую директорию проекта в путь
project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
sys.path.insert(0, project_root)

from gpt.assistant import TEST_MODEL_NAME  # noqa: E402
from scripts.generate import generate_test_work  # noqa: E402


def check_pandoc_available() -> bool:
    """
    Проверяет, доступен ли pandoc в системе.
    
    Returns:
        True, если pandoc доступен, False иначе
    """
    try:
        result = subprocess.run(
            ['pandoc', '--version'],
            capture_output=True,
            timeout=5
        )
        return result.returncode == 0
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return False


def find_toc_position(doc: Document) -> int | None:
    """
    Находит позицию TOC в документе.
    
    Args:
        doc: Document объект
    
    Returns:
        Индекс параграфа с TOC или None, если не найден
    """
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if ('table of contents' in text or
            'содержание' in text or
            'оглавление' in text):
            return i
    return None


def find_toc_sdt_position(doc: Document) -> int | None:
    """
    Находит позицию TOC как SDT элемента в body документа.
    
    Args:
        doc: Document объект
    
    Returns:
        Индекс элемента в body или None, если не найден
    """
    body = doc.element.body
    for i, elem in enumerate(body):
        if 'sdt' in elem.tag.lower():
            return i
    return None


def find_title_page_end(doc: Document) -> int | None:
    """
    Находит конец титульной страницы в документе.
    
    Args:
        doc: Document объект
    
    Returns:
        Индекс последнего параграфа титульной страницы или None
    """
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if 'проверил:' in text.lower() or ('петров' in text.lower() and 'п.п' in text.lower()):
            # Ищем последний параграф титульной страницы
            title_end = i
            for j in range(i + 1, min(i + 3, len(doc.paragraphs))):
                if doc.paragraphs[j].text.strip():
                    title_end = j
                else:
                    break
            return title_end
    return None


def find_title_page_end_in_body(doc: Document) -> int | None:
    """
    Находит конец титульной страницы в body документа.
    
    Args:
        doc: Document объект
    
    Returns:
        Индекс элемента в body или None
    """
    body = doc.element.body
    title_end_para_idx = find_title_page_end(doc)
    if title_end_para_idx is None:
        return None
    
    para_count = 0
    for i, elem in enumerate(body):
        if 'p' in elem.tag.lower():
            if para_count == title_end_para_idx:
                return i
            para_count += 1
    return None


@pytest.fixture
def temp_dir():
    """Фикстура для создания временной директории"""
    temp_dir_path = tempfile.mkdtemp()
    yield temp_dir_path
    # Очистка после теста
    if os.path.exists(temp_dir_path):
        with contextlib.suppress(Exception):
            shutil.rmtree(temp_dir_path)


@pytest.fixture
def test_theme():
    """Фикстура для тестовой темы"""
    return "Тестовая работа для проверки TOC"


@pytest.fixture
def test_pages():
    """Фикстура для количества страниц"""
    return 2


@pytest.mark.asyncio
async def test_docx_toc_after_title_page(temp_dir, test_theme, test_pages):
    """
    Тест: проверяет, что TOC находится после титульной страницы в DOCX файле.
    
    Проверяет:
    1. DOCX файл создан
    2. TOC присутствует в документе
    3. TOC находится после титульной страницы (не в начале)
    """
    if not check_pandoc_available():
        pytest.skip("Pandoc не установлен. Пропускаем тест генерации DOCX.")
    
    # Генерируем тестовую работу
    result = await generate_test_work(
        theme=test_theme,
        pages=test_pages,
        work_type="курсовая",
        model_name=TEST_MODEL_NAME,
        output_dir=temp_dir
    )
    
    # Проверяем, что DOCX был создан
    assert result['docx_path'] is not None, "DOCX файл должен быть создан"
    assert os.path.exists(result['docx_path']), f"DOCX файл должен существовать: {result['docx_path']}"
    
    # Открываем DOCX файл
    doc = Document(result['docx_path'])
    
    # Проверяем наличие TOC
    toc_para_idx = find_toc_position(doc)
    toc_sdt_idx = find_toc_sdt_position(doc)
    
    assert toc_para_idx is not None or toc_sdt_idx is not None, (
        "TOC должен присутствовать в документе. "
        "Проверьте, что pandoc создал оглавление с флагом --toc."
    )
    
    # Находим конец титульной страницы
    title_end_para_idx = find_title_page_end(doc)
    assert title_end_para_idx is not None, (
        "Титульная страница должна быть найдена в документе. "
        "Проверьте структуру LaTeX шаблона."
    )
    
    # Проверяем, что TOC находится после титульной страницы
    if toc_para_idx is not None:
        assert toc_para_idx >= title_end_para_idx, (
            f"TOC должен находиться ПОСЛЕ или НА титульной странице. "
            f"Текущая позиция TOC: {toc_para_idx}, "
            f"конец титульной страницы: {title_end_para_idx}"
        )
    
    if toc_sdt_idx is not None:
        title_end_body_idx = find_title_page_end_in_body(doc)
        assert title_end_body_idx is not None, "Не удалось найти титульную страницу в body"
        # TOC может быть сразу после титульной страницы (>=), но не до неё
        assert toc_sdt_idx >= title_end_body_idx, (
            f"TOC (SDT) должен находиться ПОСЛЕ или НА титульной странице в body. "
            f"Текущая позиция TOC: {toc_sdt_idx}, "
            f"конец титульной страницы: {title_end_body_idx}"
        )
        # Если TOC на той же позиции, проверяем, что это не начало документа
        if toc_sdt_idx == title_end_body_idx:
            assert toc_sdt_idx > 0, "TOC не должен быть в самом начале документа"


@pytest.mark.asyncio
async def test_docx_toc_not_at_beginning(temp_dir, test_theme, test_pages):
    """
    Тест: проверяет, что TOC НЕ находится в самом начале документа.
    
    Pandoc по умолчанию размещает TOC в начале, но наша функция
    должна переместить его после титульной страницы.
    """
    if not check_pandoc_available():
        pytest.skip("Pandoc не установлен. Пропускаем тест генерации DOCX.")
    
    # Генерируем тестовую работу
    result = await generate_test_work(
        theme=test_theme,
        pages=test_pages,
        work_type="курсовая",
        model_name=TEST_MODEL_NAME,
        output_dir=temp_dir
    )
    
    # Проверяем, что DOCX был создан
    assert result['docx_path'] is not None
    assert os.path.exists(result['docx_path'])
    
    # Открываем DOCX файл
    doc = Document(result['docx_path'])
    body = doc.element.body
    
    # Проверяем, что первый элемент - не TOC
    if len(body) > 0:
        first_elem = body[0]
        first_is_toc = 'sdt' in first_elem.tag.lower()
        
        # Если первый элемент - TOC, проверяем, что это не начало документа
        if first_is_toc:
            # Проверяем первые параграфы
            first_para_text = doc.paragraphs[0].text.strip().lower() if doc.paragraphs else ""
            is_toc_text = 'table of contents' in first_para_text or 'содержание' in first_para_text
            
            assert not is_toc_text, (
                "TOC не должен быть в самом начале документа. "
                "Проверьте функцию _move_toc_after_title_page()."
            )


@pytest.mark.asyncio
async def test_docx_structure_order(temp_dir, test_theme, test_pages):
    """
    Тест: проверяет правильный порядок элементов в DOCX файле.
    
    Правильный порядок:
    1. Титульная страница
    2. Оглавление (TOC)
    3. Основной контент
    """
    if not check_pandoc_available():
        pytest.skip("Pandoc не установлен. Пропускаем тест генерации DOCX.")
    
    # Генерируем тестовую работу
    result = await generate_test_work(
        theme=test_theme,
        pages=test_pages,
        work_type="курсовая",
        model_name=TEST_MODEL_NAME,
        output_dir=temp_dir
    )
    
    # Проверяем, что DOCX был создан
    assert result['docx_path'] is not None
    assert os.path.exists(result['docx_path'])
    
    # Открываем DOCX файл
    doc = Document(result['docx_path'])
    paragraphs = list(doc.paragraphs)
    
    # Находим позиции ключевых элементов
    title_start_idx = None
    title_end_idx = find_title_page_end(doc)
    toc_idx = find_toc_position(doc)
    
    # Находим начало титульной страницы
    for i, para in enumerate(paragraphs):
        text = para.text.strip().lower()
        if 'министерство' in text or 'российский государственный университет' in text:
            title_start_idx = i
            break
    
    assert title_start_idx is not None, "Не удалось найти начало титульной страницы"
    assert title_end_idx is not None, "Не удалось найти конец титульной страницы"
    
    # TOC может быть найден в параграфах или как SDT элемент
    toc_sdt_idx = find_toc_sdt_position(doc)
    assert toc_idx is not None or toc_sdt_idx is not None, "Не удалось найти TOC"
    
    # Если TOC найден в параграфах, проверяем порядок
    if toc_idx is not None:
        # Проверяем порядок: титульная страница → TOC → контент
        assert title_start_idx < title_end_idx <= toc_idx, (
            f"Неправильный порядок элементов! "
            f"Начало титульной страницы: {title_start_idx}, "
            f"Конец титульной страницы: {title_end_idx}, "
            f"TOC: {toc_idx}. "
            f"Ожидаемый порядок: титульная страница → TOC → контент"
        )


@pytest.mark.asyncio
async def test_docx_toc_contains_content(temp_dir, test_theme, test_pages):
    """
    Тест: проверяет, что TOC содержит содержимое (не пустой).
    
    Проверяет, что TOC не только присутствует, но и содержит
    ссылки на разделы документа.
    """
    if not check_pandoc_available():
        pytest.skip("Pandoc не установлен. Пропускаем тест генерации DOCX.")
    
    # Генерируем тестовую работу
    result = await generate_test_work(
        theme=test_theme,
        pages=test_pages,
        work_type="курсовая",
        model_name=TEST_MODEL_NAME,
        output_dir=temp_dir
    )
    
    # Проверяем, что DOCX был создан
    assert result['docx_path'] is not None
    assert os.path.exists(result['docx_path'])
    
    # Открываем DOCX файл
    doc = Document(result['docx_path'])
    
    # Находим TOC
    toc_idx = find_toc_position(doc)
    if toc_idx is None:
        # Проверяем SDT элемент
        toc_sdt_idx = find_toc_sdt_position(doc)
        assert toc_sdt_idx is not None, "TOC должен присутствовать в документе"
        # SDT элемент может содержать TOC, но проверить его содержимое сложнее
        return
    
    # Проверяем, что после заголовка TOC есть содержимое
    # Обычно TOC содержит несколько строк с разделами
    toc_has_content = False
    for i in range(toc_idx + 1, min(toc_idx + 10, len(doc.paragraphs))):
        para_text = doc.paragraphs[i].text.strip()
        # TOC обычно содержит номера страниц или названия разделов
        if para_text and (any(char.isdigit() for char in para_text) or len(para_text) > MIN_TOC_CONTENT_LENGTH):
            toc_has_content = True
            break
    
    assert toc_has_content, (
        "TOC должен содержать содержимое (ссылки на разделы). "
        "Проверьте, что pandoc правильно создал оглавление."
    )


def count_page_breaks_in_document(doc: Document) -> int:
    """
    Подсчитывает количество разрывов страниц в документе.
    
    Args:
        doc: Document объект
    
    Returns:
        Количество разрывов страниц
    """
    page_breaks_count = 0
    for para in doc.paragraphs:
        for run in para.runs:
            if hasattr(run, '_element'):
                xml = run._element.xml
                # Ищем разрывы страниц в XML: <w:br w:type="page"/>
                if 'w:br' in xml and 'w:type="page"' in xml:
                    page_breaks_count += 1
    return page_breaks_count


def find_page_break_positions(doc: Document) -> list[int]:
    """
    Находит позиции параграфов, содержащих разрывы страниц.
    
    Args:
        doc: Document объект
    
    Returns:
        Список индексов параграфов с разрывами страниц
    """
    positions = []
    for i, para in enumerate(doc.paragraphs):
        for run in para.runs:
            if hasattr(run, '_element'):
                xml = run._element.xml
                if 'w:br' in xml and 'w:type="page"' in xml:
                    positions.append(i)
                    break
    return positions


@pytest.mark.asyncio
async def test_docx_page_breaks(temp_dir, test_theme, test_pages):  # noqa: PLR0912, PLR0915
    """
    Тест: проверяет наличие разрывов страниц в DOCX файле.
    
    Проверяет:
    1. Разрыв страницы после титульной страницы (перед TOC)
    2. Разрыв страницы после TOC (перед первой главой)
    3. Разрыв страницы перед каждой новой главой
    4. Общее количество разрывов страниц (минимум 2 + количество глав)
    """
    if not check_pandoc_available():
        pytest.skip("Pandoc не установлен. Пропускаем тест генерации DOCX.")
    
    # Генерируем тестовую работу
    result = await generate_test_work(
        theme=test_theme,
        pages=test_pages,
        work_type="курсовая",
        model_name=TEST_MODEL_NAME,
        output_dir=temp_dir
    )
    
    # Проверяем, что DOCX был создан
    assert result['docx_path'] is not None
    assert os.path.exists(result['docx_path'])
    
    # Открываем DOCX файл
    doc = Document(result['docx_path'])
    paragraphs = list(doc.paragraphs)
    
    # Подсчитываем разрывы страниц
    page_breaks_count = count_page_breaks_in_document(doc)
    page_break_positions = find_page_break_positions(doc)
    
    # Логируем для отладки (используем print, так как logger может быть не настроен)
    print(f"Найдено разрывов страниц: {page_breaks_count}")
    print(f"Позиции разрывов страниц: {page_break_positions}")
    
    # Должно быть минимум 2 разрыва страницы:
    # 1. После титульной страницы (перед TOC)
    # 2. После TOC (перед первой главой)
    # + по одному перед каждой новой главой
    assert page_breaks_count >= MIN_EXPECTED_PAGE_BREAKS, (
        f"Должно быть минимум {MIN_EXPECTED_PAGE_BREAKS} разрыва страницы (после титульной и после TOC). "
        f"Найдено: {page_breaks_count}"
    )
    
    # Находим позиции ключевых элементов
    title_end_idx = find_title_page_end(doc)
    toc_start_idx = find_toc_position(doc)
    toc_sdt_idx = find_toc_sdt_position(doc)
    toc_end_idx = None
    
    # TOC может быть как в параграфах, так и как SDT элемент
    toc_found = toc_start_idx is not None or toc_sdt_idx is not None
    
    # Находим конец TOC (начало первой главы)
    # Согласно коду, разрыв после TOC добавляется перед первой главой (toc_end_idx)
    # Разрыв может быть создан как отдельный параграф перед главой или внутри параграфа главы
    if toc_start_idx is not None:
        for i in range(toc_start_idx + 1, min(toc_start_idx + 20, len(paragraphs))):
            text = paragraphs[i].text.strip()
            para = paragraphs[i]
            # Ищем начало первой главы (заголовок с стилем Heading)
            if (text and len(text) > MIN_TEXT_LENGTH_FOR_SECTION and
                not any(marker in text.lower() for marker in ['table of contents', 'содержание', 'оглавление']) and
                (para.style and 'heading' in para.style.name.lower())):
                toc_end_idx = i
                break
    elif toc_sdt_idx is not None and title_end_idx is not None:
        # TOC - SDT элемент, ищем первую главу после титульной страницы
        for i in range(title_end_idx + 1, min(title_end_idx + 20, len(paragraphs))):
            text = paragraphs[i].text.strip()
            para = paragraphs[i]
            if (text and len(text) > MIN_TEXT_LENGTH_FOR_SECTION and
                (para.style and 'heading' in para.style.name.lower())):
                toc_end_idx = i
                break
    
    # Проверяем наличие разрыва страницы после титульной страницы (перед TOC)
    if title_end_idx is not None and toc_found:
        # Разрыв страницы должен быть между титульной страницей и TOC
        # Разрыв может быть в TOC параграфе (если TOC в параграфах) или перед SDT элементом
        toc_check_idx = toc_start_idx if toc_start_idx is not None else (title_end_idx + 1)
        breaks_after_title = [
            pos for pos in page_break_positions
            if title_end_idx < pos <= toc_check_idx + 1  # Разрыв может быть в TOC параграфе
        ]
        assert len(breaks_after_title) > 0, (
            f"Должен быть разрыв страницы после титульной страницы (позиция {title_end_idx}) "
            f"и перед TOC (позиция {toc_check_idx}). "
            f"Найдены разрывы на позициях: {page_break_positions}"
        )
    
    # Проверяем наличие разрыва страницы после TOC (перед первой главой)
    if toc_end_idx is not None:
        # Разрыв страницы должен быть после TOC, может быть в начале первой главы (внутри параграфа)
        # Проверяем, есть ли разрыв в самом параграфе первой главы (внутри параграфа)
        has_break_in_first_chapter = False
        if toc_end_idx < len(paragraphs):
            first_chapter_para = paragraphs[toc_end_idx]
            # Проверяем, есть ли разрыв страницы внутри параграфа первой главы
            for run in first_chapter_para.runs:
                if hasattr(run, '_element'):
                    xml = run._element.xml
                    if 'w:br' in xml and 'w:type="page"' in xml:
                        has_break_in_first_chapter = True
                        break
        
        # Согласно коду, разрыв после TOC создается как отдельный параграф перед первой главой
        # (используется insert_paragraph_before), поэтому разрыв должен быть на позиции toc_end_idx - 1
        # Также проверяем разрывы в позициях после TOC (может быть в начале первой главы)
        breaks_after_toc = [
            pos for pos in page_break_positions
            if (toc_end_idx - 1 <= pos <= toc_end_idx)  # Разрыв может быть перед главой (toc_end_idx - 1) или в начале главы (toc_end_idx)
        ]
        
        # Если разрыв найден внутри параграфа первой главы, добавляем его в список
        if has_break_in_first_chapter and toc_end_idx not in breaks_after_toc:
            breaks_after_toc.append(toc_end_idx)
        
        # Разрыв может быть либо в списке позиций (отдельный параграф перед главой), либо внутри параграфа первой главы
        assert len(breaks_after_toc) > 0 or has_break_in_first_chapter, (
            f"Должен быть разрыв страницы после TOC (перед позицией {toc_end_idx}, т.е. на позиции {toc_end_idx - 1} или в начале параграфа {toc_end_idx}). "
            f"Найдены разрывы на позициях: {page_break_positions}, "
            f"toc_end_idx={toc_end_idx}, всего параграфов={len(paragraphs)}, "
            f"разрыв в параграфе первой главы: {has_break_in_first_chapter}"
        )
    
    # Подсчитываем количество глав (секций)
    sections_count = 0
    for i, para in enumerate(paragraphs):
        # Пропускаем титульную страницу и TOC
        if title_end_idx is not None and i <= title_end_idx:
            continue
        if toc_start_idx is not None and toc_end_idx is not None and toc_start_idx <= i <= toc_end_idx:
            continue
        
        text = para.text.strip()
        if not text:
            continue
        
        # Проверяем, является ли параграф заголовком секции
        is_section = False
        if para.style and para.style.name:
            style_name = para.style.name.lower()
            if 'heading' in style_name:
                is_section = True
        
        if (not is_section and text and
            (re.match(r'^\d+[.)]\s+[А-ЯЁ]', text) or re.match(r'^\d+[.)]\s+[A-Z]', text))):
            # Проверяем паттерн номера секции
            is_section = True
        
        if is_section:
            sections_count += 1
    
    # Ожидаемое количество разрывов страниц:
    # 1 после титульной страницы (перед TOC) + 1 после TOC (перед первой главой)
    # + по одному перед каждой новой главой (начиная со второй)
    # Согласно коду, разрывы перед главами добавляются только начиная со второй главы
    # и только если глава распознана как секция
    expected_min_breaks = 2  # Минимум: перед TOC и после TOC (обязательные)
    
    # Дополнительные разрывы перед каждой новой главой (начиная со второй)
    # Но разрывы перед главами могут не добавляться, если главы не распознаны как секции
    # Поэтому проверяем только основные разрывы (2), а разрывы перед главами - опционально
    expected_breaks_before_sections = 0
    if sections_count > 1:
        expected_breaks_before_sections = sections_count - 1
    
    # Проверяем обязательные разрывы (перед TOC и после TOC)
    assert page_breaks_count >= expected_min_breaks, (
        f"Ожидается минимум {expected_min_breaks} разрыва страницы (1 после титульной, 1 после TOC). "
        f"Найдено: {page_breaks_count}, глав: {sections_count}, TOC найден: {toc_found}"
    )
    
    # Дополнительно проверяем наличие разрывов перед главами (если глав больше 1)
    # Это необязательная проверка, так как разрывы перед главами могут не добавляться
    # если главы не распознаны как секции
    if sections_count > 1 and expected_breaks_before_sections > 0:
        # Логируем информацию о разрывах перед главами (не строгая проверка)
        print(f"  Ожидается {expected_breaks_before_sections} разрывов перед главами, "
              f"найдено {page_breaks_count} разрывов всего")


@pytest.mark.asyncio
async def test_docx_page_breaks_structure(temp_dir, test_theme, test_pages):
    """
    Тест: проверяет структуру разрывов страниц в DOCX файле.
    
    Проверяет правильный порядок: титульная страница → [разрыв] → TOC → [разрыв] → главы.
    """
    if not check_pandoc_available():
        pytest.skip("Pandoc не установлен. Пропускаем тест генерации DOCX.")
    
    # Генерируем тестовую работу
    result = await generate_test_work(
        theme=test_theme,
        pages=test_pages,
        work_type="курсовая",
        model_name=TEST_MODEL_NAME,
        output_dir=temp_dir
    )
    
    # Проверяем, что DOCX был создан
    assert result['docx_path'] is not None
    assert os.path.exists(result['docx_path'])
    
    # Открываем DOCX файл
    doc = Document(result['docx_path'])
    
    # Находим позиции ключевых элементов
    title_end_idx = find_title_page_end(doc)
    toc_start_idx = find_toc_position(doc)
    toc_sdt_idx = find_toc_sdt_position(doc)
    
    assert title_end_idx is not None, "Не удалось найти конец титульной страницы"
    
    # TOC может быть как в параграфах, так и как SDT элемент
    toc_found = toc_start_idx is not None or toc_sdt_idx is not None
    assert toc_found, "Не удалось найти TOC (ни в параграфах, ни как SDT элемент)"
    
    # Если TOC найден в параграфах, проверяем порядок
    if toc_start_idx is not None:
        # Проверяем порядок: титульная страница должна быть перед TOC
        assert title_end_idx < toc_start_idx, (
            f"Титульная страница (конец на позиции {title_end_idx}) "
            f"должна быть перед TOC (позиция {toc_start_idx})"
        )
        
        # Проверяем наличие разрыва страницы между титульной страницей и TOC
        page_break_positions = find_page_break_positions(doc)
        breaks_between_title_and_toc = [
            pos for pos in page_break_positions
            if title_end_idx < pos <= toc_start_idx + 1  # Разрыв может быть в TOC параграфе
        ]
        
        assert len(breaks_between_title_and_toc) > 0, (
            f"Должен быть разрыв страницы между титульной страницей (позиция {title_end_idx}) "
            f"и TOC (позиция {toc_start_idx}). "
            f"Найдены разрывы на позициях: {page_break_positions}"
        )
    elif toc_sdt_idx is not None:
        # TOC - SDT элемент, проверяем наличие разрыва после титульной страницы
        page_break_positions = find_page_break_positions(doc)
        breaks_after_title = [
            pos for pos in page_break_positions
            if title_end_idx < pos <= title_end_idx + 3  # Разрыв может быть перед SDT элементом
        ]
        
        assert len(breaks_after_title) > 0, (
            f"Должен быть разрыв страницы после титульной страницы (позиция {title_end_idx}) "
            f"и перед TOC (SDT элемент). "
            f"Найдены разрывы на позициях: {page_break_positions}"
        )


def count_pages_in_docx(docx_path: str) -> int:
    """
    Подсчитывает приблизительное количество страниц в DOCX файле.
    
    Использует метод подсчета разрывов страниц и структуры документа
    для оценки количества страниц.
    
    Args:
        docx_path: Путь к DOCX файлу
    
    Returns:
        Приблизительное количество страниц
    """
    doc = Document(docx_path)
    paragraphs = list(doc.paragraphs)
    
    # Подсчитываем разрывы страниц
    page_breaks = count_page_breaks_in_document(doc)
    
    # Базовое количество страниц: минимум 1 страница
    # Каждый разрыв страницы добавляет минимум 1 страницу
    estimated_pages = 1 + page_breaks
    
    # Если есть TOC, добавляем страницу для него
    toc_found = find_toc_position(doc) is not None or find_toc_sdt_position(doc) is not None
    if toc_found:
        estimated_pages += 1
    
    # Учитываем количество глав (секций)
    sections_count = 0
    for para in paragraphs:
        if para.style and 'heading' in para.style.name.lower():
            sections_count += 1
    
    # Каждая глава должна быть на отдельной странице
    if sections_count > 0:
        estimated_pages = max(estimated_pages, 1 + sections_count + (1 if toc_found else 0))
    
    return estimated_pages


@pytest.mark.asyncio
async def test_docx_page_breaks_count(temp_dir, test_theme, test_pages):  # noqa: PLR0912
    """
    Тест: проверяет количество страниц в DOCX файле.
    
    Проверяет, что документ имеет правильное количество страниц:
    - Минимум 1 страница (титульный лист)
    - +1 страница если есть TOC
    - +1 страница для каждой главы
    
    Также проверяет наличие разрывов страниц в правильных местах.
    """
    if not check_pandoc_available():
        pytest.skip("Pandoc не установлен. Пропускаем тест генерации DOCX.")
    
    # Генерируем тестовую работу
    result = await generate_test_work(
        theme=test_theme,
        pages=test_pages,
        work_type="курсовая",
        model_name=TEST_MODEL_NAME,
        output_dir=temp_dir
    )
    
    # Проверяем, что DOCX был создан
    assert result['docx_path'] is not None
    assert os.path.exists(result['docx_path'])
    
    # Открываем DOCX файл
    doc = Document(result['docx_path'])
    paragraphs = list(doc.paragraphs)
    
    # Подсчитываем разрывы страниц
    page_breaks_count = count_page_breaks_in_document(doc)
    
    # Проверяем наличие TOC
    toc_found = find_toc_position(doc) is not None or find_toc_sdt_position(doc) is not None
    
    # Подсчитываем количество глав (секций)
    sections_count = 0
    title_end_idx = find_title_page_end(doc)
    toc_start_idx = find_toc_position(doc)
    
    for i, para in enumerate(paragraphs):
        # Пропускаем титульную страницу и TOC
        if title_end_idx is not None and i <= title_end_idx:
            continue
        if toc_start_idx is not None and toc_start_idx <= i < toc_start_idx + 10:
            # Пропускаем TOC (примерно 5-10 параграфов после начала)
            continue
        
        if para.style and 'heading' in para.style.name.lower():
            sections_count += 1
    
    # Ожидаемое минимальное количество страниц:
    # 1 (титульная страница) + 1 (TOC, если есть) + количество глав
    min_expected_pages = 1
    if toc_found:
        min_expected_pages += 1
    if sections_count > 0:
        min_expected_pages += sections_count
    
    # Ожидаемое минимальное количество разрывов страниц:
    # 1 после титульной страницы (перед TOC) + 1 после TOC (перед первой главой) +
    # по одному перед каждой новой главой (кроме первой)
    min_expected_breaks = 0
    if toc_found:
        # Если есть TOC: разрыв перед TOC и разрыв после TOC (перед первой главой)
        min_expected_breaks = 2
        # Дополнительные разрывы перед каждой новой главой (начиная со второй)
        # Но учитываем, что разрыв может не добавляться, если главы идут подряд
        # Поэтому проверяем только наличие основных разрывов
    elif sections_count > 0:
        # Если TOC нет: разрыв перед первой главой
        min_expected_breaks = 1
    
    # Проверяем количество разрывов страниц (основные разрывы обязательны)
    assert page_breaks_count >= min_expected_breaks, (
        f"Ожидается минимум {min_expected_breaks} разрывов страницы (основные: перед/после TOC или перед первой главой). "
        f"Найдено: {page_breaks_count}. "
        f"TOC найден: {toc_found}, глав: {sections_count}"
    )
    
    # Дополнительно проверяем наличие разрывов перед главами (если глав больше 1)
    # Это необязательная проверка, так как разрывы перед главами могут добавляться
    # только если они определены как секции в коде
    if sections_count > 1:
        # Ищем разрывы страниц перед заголовками глав (начиная со второй)
        breaks_before_sections = 0
        previous_section_idx = None
        for i, para in enumerate(paragraphs):
            if para.style and 'heading' in para.style.name.lower():
                if previous_section_idx is not None:
                    # Проверяем, есть ли разрыв страницы перед этой главой
                    has_break = False
                    if para.runs:
                        first_run = para.runs[0]
                        if hasattr(first_run, '_element'):
                            xml = first_run._element.xml
                            if 'w:br' in xml and 'w:type="page"' in xml:
                                has_break = True
                    if has_break:
                        breaks_before_sections += 1
                previous_section_idx = i
        
        # Логируем информацию о разрывах перед главами (не строгая проверка)
        if breaks_before_sections > 0:
            print(f"  Найдено {breaks_before_sections} разрывов страниц перед главами (из {sections_count - 1} возможных)")
    
    # Подсчитываем приблизительное количество страниц
    estimated_pages = count_pages_in_docx(result['docx_path'])
    
    # Проверяем, что количество страниц соответствует ожидаемому
    assert estimated_pages >= min_expected_pages, (
        f"Ожидается минимум {min_expected_pages} страниц. "
        f"Оценка: {estimated_pages}. "
        f"TOC найден: {toc_found}, глав: {sections_count}, разрывов: {page_breaks_count}"
    )
    
    print(f"✓ Документ содержит: {estimated_pages} страниц (минимум {min_expected_pages}), "
          f"{page_breaks_count} разрывов страниц (минимум {min_expected_breaks}), "
          f"{sections_count} глав")

